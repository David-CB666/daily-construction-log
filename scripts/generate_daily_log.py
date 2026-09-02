#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
每日施工記錄表 — 一鍵生成腳本 v3.0
=====================================
用途：從昨日docx複製 → 修改變動欄位 → 嵌入照片 → 輸出docx
項目：[項目名稱]
更新：2026-08-01（照片規格寫死2.6"×2.0"、支援多種照片命名）

使用方法：
    python generate_daily_log.py --date 2026-08-02 --workers 4 --tasks "工序1,工序2,工序3"

環境需求：
    pip install python-docx docx2pdf pypdf pymupdf
"""

import os
import sys
import shutil
import argparse
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================
# 🔧 配置區 — 按項目修改
# ============================================================
BASE_DIR = r"<PROJECTS_ROOT>\[項目名稱]\每日施工記錄"
PROJECT_START_DATE = "2026-06-22"  # 開工第一天
DEFAULT_ENTRY_TIME = "09:00"
DEFAULT_EXIT_TIME = "18:00"
RECORDER = "[建築公司名稱]"

# ============================================================
# 📸 照片規格（🔴 不可修改 — Mike 2026-07-31 確立）
# ============================================================
PHOTO_WIDTH_INCH = 2.6
PHOTO_HEIGHT_INCH = 2.0

# ============================================================
# 📋 天氣數據提取函數
# ============================================================
def extract_weather(date_str):
    """
    從 SMG 天氣報告 PDF 提取天氣描述。
    優先查找附件/天氣報告/天氣報告_YYYY-MM-DD.pdf
    返回: 天氣描述文字
    """
    import fitz
    weather_pdf = os.path.join(BASE_DIR, "附件", "天氣報告", f"天氣報告_{date_str}.pdf")
    if not os.path.exists(weather_pdf):
        print(f"⚠️ 天氣報告不存在: {weather_pdf}")
        return "（待填寫）"

    doc = fitz.open(weather_pdf)
    text = ""
    for page in doc:
        text += page.get_text()

    # 嘗試提取中文天氣描述
    lines = text.split('\n')
    chinese_lines = []
    in_chinese_section = False
    for line in lines:
        if '今日' in line and '天氣預報' in line:
            in_chinese_section = True
            continue
        if in_chinese_section:
            if 'PREVISÃO' in line or 'Vento' in line or 'BM4' in line:
                break
            stripped = line.strip()
            if stripped and not stripped.startswith('http') and '[地區]' not in stripped:
                chinese_lines.append(stripped)

    if chinese_lines:
        weather = "".join(chinese_lines).replace("。", "。").strip()
        if weather.endswith("。"):
            pass
        return weather

    return "（待填寫）"


def extract_tide(date_str):
    """
    從 SMG 潮汐預報 PDF 提取潮汐數據。
    使用圖像渲染 + AI 辨識方式。
    返回: (low_tide_str, high_tide_str) 如 ("1.0m（約17:00）", "2.8m（約08:00）")
    """
    tide_pdf = os.path.join(BASE_DIR, "附件", "潮汐預報", f"潮汐預報_{date_str}.pdf")
    if not os.path.exists(tide_pdf):
        print(f"⚠️ 潮汐預報不存在: {tide_pdf}")
        return "（待填寫）", "（待填寫）"

    import fitz
    doc = fitz.open(tide_pdf)
    page = doc[0]
    pix = page.get_pixmap(dpi=200)

    # 暫存PNG
    png_path = os.path.join(BASE_DIR, f"_tide_{date_str}.png")
    pix.save(png_path)
    print(f"📊 潮汐預報已渲染: {png_path}")
    print("⚠️ 請使用AI圖像辨識讀取潮汐數據，然後手動填入")
    return "（待AI辨識）", "（待AI辨識）"


# ============================================================
# 📸 照片掃描與嵌入
# ============================================================
def find_photos(date_str):
    """
    掃描照片目錄，支援多種命名格式。
    返回: 照片路徑列表（最多4張）
    """
    photo_dir = os.path.join(BASE_DIR, "附件", "現埸相", date_str)
    if not os.path.isdir(photo_dir):
        print(f"⚠️ 照片目錄不存在: {photo_dir}")
        return []

    # 支援的命名格式
    naming_schemes = [
        ["圖一.jpg", "圖二.jpg", "圖三.jpg", "圖四.jpg"],
        ["圖 (1).jpg", "圖 (2).jpg", "圖 (3).jpg", "圖 (4).jpg"],
        ["圖1.jpg", "圖2.jpg", "圖3.jpg", "圖4.jpg"],
    ]

    photos = []
    for scheme in naming_schemes:
        photos = []
        for fname in scheme:
            fpath = os.path.join(photo_dir, fname)
            if os.path.exists(fpath):
                photos.append(fpath)
        if len(photos) >= 1:
            break

    # 如果命名格式都不匹配，取目錄中最先4個jpg文件
    if not photos:
        jpgs = sorted([f for f in os.listdir(photo_dir) if f.lower().endswith('.jpg')])
        for fname in jpgs[:4]:
            photos.append(os.path.join(photo_dir, fname))

    return photos


def clear_table5_photos(table5):
    """清除表格5所有儲存格內的舊內容"""
    for row_idx in range(len(table5.rows)):
        for col_idx in range(len(table5.rows[row_idx].cells)):
            cell = table5.rows[row_idx].cells[col_idx]
            for p in cell.paragraphs:
                p.clear()


def embed_photos(table5, photos):
    """
    嵌入照片到表格5。
    圖一→R0C0, 圖二→R0C1, 圖三→R1C0, 圖四→R1C1
    規格：2.6" × 2.0"（🔴 不可修改）
    """
    cell_map = [(0, 0), (0, 1), (1, 0), (1, 1)]
    for idx, (row_idx, col_idx) in enumerate(cell_map):
        if idx < len(photos):
            cell = table5.rows[row_idx].cells[col_idx]
            p = cell.paragraphs[0]
            run = p.add_run()
            run.add_picture(photos[idx], width=Inches(PHOTO_WIDTH_INCH), height=Inches(PHOTO_HEIGHT_INCH))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ============================================================
# 🧮 天數計算
# ============================================================
def calc_day_number(date_str):
    """計算開工第幾天（6/22 = 第1天）"""
    from datetime import date
    start = date.fromisoformat(PROJECT_START_DATE)
    today = date.fromisoformat(date_str)
    return (today - start).days + 1


def find_yesterday_docx(date_str):
    """找到最近一日的記錄表docx"""
    from datetime import date, timedelta
    today = date.fromisoformat(date_str)
    yesterday = today - timedelta(days=1)

    # 嘗試昨日、前日、大前日...
    for offset in range(1, 8):
        d = today - timedelta(days=offset)
        d_str = d.isoformat()
        # 支援多種命名格式
        for fmt in [f"{d_str}每日施工記錄表.docx", f"{d_str}_每日施工記錄表.docx"]:
            path = os.path.join(BASE_DIR, fmt)
            if os.path.exists(path):
                return path

    raise FileNotFoundError(f"找不到最近一日的記錄表（回溯7天）")


# ============================================================
# 🔧 主生成邏輯
# ============================================================
def generate(date_str, workers, tasks, stage=None, weather_override=None, tide_low_override=None, tide_high_override=None):
    """
    生成每日施工記錄表。

    參數：
        date_str:       日期 YYYY-MM-DD
        workers:        施工人數（如 "4人"）
        tasks:          施工項目列表 [(序號, 項目, 狀態), ...]
        stage:          施工階段（如不提供則繼承昨日）
        weather_override: 手動天氣覆蓋（如不提供則自動提取）
        tide_low_override: 手動低潮覆蓋
        tide_high_override: 手動高潮覆蓋
    """
    day_num = calc_day_number(date_str)

    # 1. 複製昨日文檔
    yesterday_path = find_yesterday_docx(date_str)
    today_path = os.path.join(BASE_DIR, f"{date_str}每日施工記錄表.docx")
    shutil.copy(yesterday_path, today_path)
    print(f"📋 已複製: {os.path.basename(yesterday_path)} → {os.path.basename(today_path)}")

    doc = Document(today_path)
    tables = doc.tables

    # 2. 修改 Table 0 — 基本信息
    t0 = tables[0]
    p = t0.rows[0].cells[1].paragraphs[0]
    p.clear()
    p.add_run(date_str)

    p = t0.rows[0].cells[3].paragraphs[0]
    p.clear()
    p.add_run(f"第{day_num}天")

    # 施工階段：繼承昨日前綴或使用覆蓋值
    if stage:
        p = t0.rows[1].cells[1].paragraphs[0]
        p.clear()
        p.add_run(stage)
    # 否則不修改，繼承昨日

    # 3. 修改 Table 1 — 天氣與環境
    t1 = tables[1]
    # 天氣
    weather = weather_override or extract_weather(date_str)
    p = t1.rows[1].cells[0].paragraphs[0]
    p.clear()
    p.add_run(weather)

    # 潮汐
    tide_low = tide_low_override
    tide_high = tide_high_override
    if not tide_low or not tide_high:
        low, high = extract_tide(date_str)
        tide_low = tide_low or low
        tide_high = tide_high or high

    p = t1.rows[1].cells[1].paragraphs[0]
    p.clear()
    p.add_run(tide_low)

    p = t1.rows[1].cells[2].paragraphs[0]
    p.clear()
    p.add_run(tide_high)

    # 警戒潮位判斷
    # 嘗試從高潮字串中提取數字
    import re
    high_match = re.search(r'([\d.]+)\s*m', tide_high or "")
    is_warning = False
    if high_match:
        high_val = float(high_match.group(1))
        is_warning = high_val >= 3.0

    cell_warn = t1.rows[1].cells[3]
    for p in cell_warn.paragraphs:
        p.clear()
    p = cell_warn.paragraphs[0]
    warning_text = "☑  是  ☐  否" if is_warning else "☐  是  ☑  否"
    run = p.add_run(warning_text)
    run.font.size = Pt(9)

    # 4. 修改 Table 2 — 施工人員
    t2 = tables[2]
    p = t2.rows[1].cells[0].paragraphs[0]
    p.clear()
    p.add_run(workers)

    # 5. 修改 Table 3 — 施工內容
    t3 = tables[3]
    # 確保有足夠行數
    num_items = len(tasks)
    current_rows = len(t3.rows) - 1  # 減去表頭
    # 若需要更多行，在這裡不處理（本版本假設行數一致）

    for idx, (num, task, status) in enumerate(tasks):
        if idx + 1 < len(t3.rows):
            row = t3.rows[idx + 1]
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.clear()
            row.cells[0].paragraphs[0].add_run(str(num))
            row.cells[1].paragraphs[0].add_run(task)
            row.cells[2].paragraphs[0].add_run(status)

    # 6. Table 4 — 安全自檢：保持不變

    # 7. Table 5 — 照片
    t5 = tables[5]
    clear_table5_photos(t5)
    photos = find_photos(date_str)
    if photos:
        embed_photos(t5, photos)
        print(f"📸 已嵌入 {len(photos)} 張照片")

    # 8. 保存
    doc.save(today_path)
    print(f"✅ 已保存: {today_path}")

    return today_path


def convert_to_pdf_and_merge(date_str):
    """將 docx 轉 PDF 並合併天氣+潮汐報告"""
    from docx2pdf import convert
    from pypdf import PdfWriter, PdfReader

    docx_path = os.path.join(BASE_DIR, f"{date_str}每日施工記錄表.docx")
    temp_pdf = os.path.join(BASE_DIR, f"{date_str}_temp.pdf")
    final_pdf = os.path.join(BASE_DIR, f"{date_str}每日施工記錄表.pdf")
    weather_pdf = os.path.join(BASE_DIR, "附件", "天氣報告", f"天氣報告_{date_str}.pdf")
    tide_pdf = os.path.join(BASE_DIR, "附件", "潮汐預報", f"潮汐預報_{date_str}.pdf")

    # docx → PDF
    print("🔄 轉換 docx → PDF...")
    convert(docx_path, temp_pdf)

    # 合併
    writer = PdfWriter()
    for f in [temp_pdf, weather_pdf, tide_pdf]:
        if os.path.exists(f):
            reader = PdfReader(f)
            for page in reader.pages:
                writer.add_page(page)
        else:
            print(f"⚠️ 跳過不存在: {f}")

    with open(final_pdf, 'wb') as f:
        writer.write(f)

    # 清理
    if os.path.exists(temp_pdf):
        os.remove(temp_pdf)

    print(f"✅ PDF 已合併: {final_pdf}")


# ============================================================
# CLI 入口
# ============================================================
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="每日施工記錄表生成器 v3.0")
    parser.add_argument("--date", required=True, help="日期 YYYY-MM-DD")
    parser.add_argument("--workers", required=True, help="施工人數（如 '4人'）")
    parser.add_argument("--tasks", required=True, help="施工項目，逗號分隔（如 '安裝圍網,安裝水電樁,堤岸燈安裝'）")
    parser.add_argument("--stage", default=None, help="施工階段覆蓋（如不提供則繼承昨日）")
    parser.add_argument("--weather", default=None, help="手動天氣描述（如不提供則自動提取）")
    parser.add_argument("--tide-low", default=None, help="手動低潮（如 '1.0m（約17:00）'）")
    parser.add_argument("--tide-high", default=None, help="手動高潮（如 '2.8m（約08:00）'）")
    parser.add_argument("--no-pdf", action="store_true", help="不生成合併PDF")

    args = parser.parse_args()

    # 解析任務
    task_list = []
    task_items = [t.strip() for t in args.tasks.split(",")]
    for i, item in enumerate(task_items, 1):
        task_list.append((str(i), item, "進行中"))

    # 生成
    generate(
        date_str=args.date,
        workers=args.workers,
        tasks=task_list,
        stage=args.stage,
        weather_override=args.weather,
        tide_low_override=args.tide_low,
        tide_high_override=args.tide_high,
    )

    # PDF
    if not args.no_pdf:
        convert_to_pdf_and_merge(args.date)
